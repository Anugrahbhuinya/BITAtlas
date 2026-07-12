import logging
import io
import docx

logger = logging.getLogger("docx_processor")

def process_docx(file_content: bytes, filename: str) -> str:
    """
    Parses a DOCX file using python-docx and extracts clean text.
    Handles headings and paragraphs.
    """
    try:
        doc = docx.Document(io.BytesIO(file_content))
        lines = []
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            # If paragraph has heading style, format it specifically
            if p.style and p.style.name and p.style.name.startswith("Heading"):
                lines.append(f"\n{text}\n")
            else:
                lines.append(text)
                
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    lines.append(row_text)
                    
        cleaned_text = "\n".join(line for line in lines if line)
        return cleaned_text
    except Exception as e:
        logger.error(f"Failed to process DOCX file {filename}: {e}", exc_info=True)
        raise ValueError(f"Failed to read DOCX document: {str(e)}")
