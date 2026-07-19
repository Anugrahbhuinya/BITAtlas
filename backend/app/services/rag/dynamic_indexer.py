import os
import uuid
import hashlib
import logging
import json
from datetime import datetime, timezone
from typing import List, Tuple, AsyncGenerator, Dict, Any
from pypdf import PdfReader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.services.rag.vector_store import get_vector_store, PERSIST_DIRECTORY
from app.core.database import get_database
from app.services.ai.cache import clear_response_cache
from app.services.rag.retriever import clear_retriever_cache

logger = logging.getLogger("dynamic_indexer")

# Absolute path to upload directory
BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
UPLOAD_DIR = os.path.join(BACKEND_DIR, "uploads", "pdfs")

def calculate_sha256(content: bytes) -> str:
    """Calculate SHA256 hash of file content."""
    sha256_hash = hashlib.sha256()
    sha256_hash.update(content)
    return sha256_hash.hexdigest()

def extract_pdf_text_by_page(file_path: str) -> List[Tuple[str, int]]:
    """
    Extracts text page-by-page from a PDF file.
    Returns a list of (page_text, page_number_1_indexed).
    """
    pages_data = []
    reader = PdfReader(file_path)
    
    if reader.is_encrypted:
        raise ValueError("PDF is encrypted. Decrypt the PDF before uploading.")
        
    for idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_data.append((text.strip(), idx + 1))
        
    return pages_data

def extract_text_from_doc(file_path: str) -> str:
    """
    Extracts text from legacy Word (.doc) binary files.
    """
    # 1. Try win32com as primary on Windows
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = None
        try:
            doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
            text = doc.Content.Text
            return text
        finally:
            if doc:
                doc.Close(False)
            word.Quit()
    except Exception as e:
        logger.warning(f"win32com extraction failed or Word not installed: {e}. Falling back to binary extraction.")
        
    # 2. Pure-Python Fallback (Extracts ASCII and UTF-16 strings from the binary structure)
    try:
        with open(file_path, "rb") as f:
            data = f.read()
            
        import re
        # Find UTF-16LE printable characters: (printable ASCII + \x00)
        # Printable ASCII is 0x20-0x7E, plus 0x09 (tab), 0x0A (LF), 0x0D (CR).
        utf16_pattern = re.compile(rb'(?:[\x20-\x7e\x09\x0a\x0d]\x00){4,}')
        ascii_pattern = re.compile(rb'[\x20-\x7e\x09\x0a\x0d]{4,}')
        
        extracted_parts = []
        # Find all UTF-16LE text blocks
        for match in utf16_pattern.finditer(data):
            try:
                decoded = match.group(0).decode("utf-16le")
                if len(decoded.strip()) > 3:
                    extracted_parts.append(decoded.strip())
            except Exception:
                pass
                
        # If UTF-16 extraction yields little text, fall back to ASCII extraction
        if not extracted_parts or len(" ".join(extracted_parts)) < 50:
            for match in ascii_pattern.finditer(data):
                try:
                    decoded = match.group(0).decode("ascii")
                    if len(decoded.strip()) > 3:
                        extracted_parts.append(decoded.strip())
                except Exception:
                    pass
                    
        # Join extracted segments
        full_text = "\n\n".join(extracted_parts)
        
        # Simple cleanup of OLE headers/footers if matched
        cleaned_lines = []
        for line in full_text.splitlines():
            if any(term in line for term in ["Microsoft Word", "MSWordDoc", "Word.Document"]):
                if len(line) < 100:
                    continue
            cleaned_lines.append(line)
            
        return "\n".join(cleaned_lines)
    except Exception as ex:
        raise ValueError(f"Failed to extract text from Word document: {str(ex)}")

def extract_txt_text(file_path: str) -> str:
    """
    Decodes txt or md file with standard encoding fallbacks.
    """
    with open(file_path, "rb") as f:
        content = f.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode text file. Supported encodings: UTF-8, UTF-16, Latin-1.")

def extract_text_by_page(file_path: str) -> List[Tuple[str, int]]:
    """
    Extracts text page-by-page from the supported document types.
    For non-paginated formats, the entire text is returned as page 1.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_pdf_text_by_page(file_path)
        
    elif ext == ".docx":
        import docx
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            text = "\n".join(full_text)
            return [(text, 1)]
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
            
    elif ext == ".doc":
        text = extract_text_from_doc(file_path)
        return [(text, 1)]
        
    elif ext in (".txt", ".md"):
        text = extract_txt_text(file_path)
        return [(text, 1)]
        
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def chunk_pdf_pages(pages_data: List[Tuple[str, int]], source_name: str, doc_id: str, file_hash: str) -> List[Document]:
    """
    Chunks each page's text using RecursiveCharacterTextSplitter (chunk_size=500, chunk_overlap=100)
    preserving page numbers. Works with multiple document formats.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )
    
    # Dynamically extract extension for metadata type fields
    ext = os.path.splitext(source_name)[1].lower()
    file_type = ext[1:] if ext else "pdf"
    
    documents = []
    chunk_idx = 0
    for text, page_num in pages_data:
        if not text:
            continue
        chunks = text_splitter.split_text(text)
        for chunk in chunks:
            # Set metadata standard keys + keep backward compatible ones
            doc = Document(
                page_content=chunk,
                metadata={
                    "document_id": doc_id,
                    "filename": source_name,
                    "page": page_num,
                    "source": "kb_document",
                    "source_type": file_type,
                    "chunk_number": chunk_idx,
                    "sha256": file_hash,
                    # Backwards compatibility keys
                    "name": source_name,
                    "title": source_name,
                    "doc_id": doc_id,
                    "type": file_type
                }
            )
            documents.append(doc)
            chunk_idx += 1
            
    return documents

async def index_pdf_generator(
    file_content: bytes, 
    filename: str, 
    overwrite: bool = False
) -> AsyncGenerator[str, None]:
    """
    Validates, saves, extracts, chunks, embeds, and indexes a supported document format.
    Yields progress status as line-delimited JSON.
    Supports atomic rollback in case of failure.
    """
    db = get_database()
    vector_store = get_vector_store()
    
    # 1. Start Uploading
    yield json.dumps({"status": "Uploading", "progress": 10}) + "\n"
    
    ext = os.path.splitext(filename)[1].lower()
    
    # Validation based on format
    if ext == ".pdf":
        if not file_content.startswith(b"%PDF"):
            yield json.dumps({"status": "Failed", "detail": "Invalid file format. Only PDF files are supported."}) + "\n"
            return
    elif ext == ".docx":
        if not file_content.startswith(b"PK\x03\x04"):
            yield json.dumps({"status": "Failed", "detail": "Invalid file format. Only DOCX files are supported."}) + "\n"
            return
    elif ext == ".doc":
        if not file_content.startswith(b"\xd0\xcf\x11\xe0"):
            yield json.dumps({"status": "Failed", "detail": "Invalid file format. Only DOC files are supported."}) + "\n"
            return
    elif ext in (".txt", ".md"):
        if b"\x00" in file_content[:1024]:
            yield json.dumps({"status": "Failed", "detail": "Binary characters detected in text file."}) + "\n"
            return
    else:
        yield json.dumps({"status": "Failed", "detail": f"Unsupported file format: {ext}"}) + "\n"
        return
        
    if len(file_content) > 10 * 1024 * 1024: # 10MB limit
        yield json.dumps({"status": "Failed", "detail": "File size exceeds 10MB limit."}) + "\n"
        return
        
    # Calculate SHA256 and check for duplicates
    file_hash = calculate_sha256(file_content)
    existing_doc = await db.indexed_documents.find_one({"hash": file_hash})
    
    if existing_doc:
        if not overwrite:
            yield json.dumps({
                "status": "Duplicate", 
                "detail": f"A document with the same content already exists: {existing_doc['filename']}",
                "doc_id": existing_doc["id"]
            }) + "\n"
            return
        else:
            # Overwrite requested. Delete the existing document first.
            yield json.dumps({"status": "Uploading", "detail": f"Overwriting existing document: {existing_doc['filename']}..."}) + "\n"
            try:
                await delete_indexed_document(existing_doc["id"])
            except Exception as e:
                logger.error(f"Error during overwrite deletion of {existing_doc['id']}: {e}")
                
    # Initialize variables for atomic rollback tracking
    doc_id = f"doc_{uuid.uuid4()}"
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(UPLOAD_DIR, f"{doc_id}{ext}")
    
    saved_on_disk = False
    vector_ids: List[str] = []
    metadata_saved = False
    
    try:
        # Write file to disk
        with open(file_path, "wb") as f:
            f.write(file_content)
        saved_on_disk = True
        
        # 2. Extracting text
        yield json.dumps({"status": "Extracting", "progress": 30}) + "\n"
        try:
            pages_data = extract_text_by_page(file_path)
        except ValueError as e:
            yield json.dumps({"status": "Failed", "detail": str(e)}) + "\n"
            # Cleanup immediately
            if os.path.exists(file_path):
                os.remove(file_path)
            return
            
        # 3. Chunking
        yield json.dumps({"status": "Chunking", "progress": 50}) + "\n"
        chunks = chunk_pdf_pages(pages_data, filename, doc_id, file_hash)
        
        if not chunks:
            yield json.dumps({"status": "Failed", "detail": "The document contains no extractable text."}) + "\n"
            if os.path.exists(file_path):
                os.remove(file_path)
            return
            
        # 4. Embedding
        yield json.dumps({"status": "Embedding", "progress": 70}) + "\n"
        # LangChain Chroma handles embedding internally, but we prepare vector IDs for tracking.
        vector_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
        
        # 5. Updating Chroma
        yield json.dumps({"status": "Updating Chroma", "progress": 85}) + "\n"
        # Sync-like execution block wrapped in try for ChromaDB insertion
        vector_store.add_documents(chunks, ids=vector_ids)
        
        # 6. Saving Metadata
        yield json.dumps({"status": "Saving Metadata", "progress": 95}) + "\n"
        metadata = {
            "id": doc_id,
            "filename": filename,
            "hash": file_hash,
            "filepath": file_path,
            "type": ext[1:],
            "status": "Indexed",
            "created": datetime.now(timezone.utc),
            "size_bytes": len(file_content),
            "chunks_count": len(chunks),
            "vector_ids": vector_ids
        }
        await db.indexed_documents.insert_one(metadata)
        metadata_saved = True
        clear_response_cache()
        clear_retriever_cache()
        
        # Log Admin Activity
        from app.services.admin_service import log_admin_activity
        await log_admin_activity(
            action="Document Uploaded & Indexed",
            username="admin",  # Can be overridden or passed down
            details={"filename": filename, "doc_id": doc_id, "chunks": len(chunks)}
        )
        
        # 7. Completed
        yield json.dumps({
            "status": "Completed", 
            "progress": 100, 
            "doc_id": doc_id,
            "filename": filename,
            "size_bytes": len(file_content),
            "chunks_count": len(chunks)
        }) + "\n"
        
    except Exception as e:
        logger.error(f"Failed to index document {filename}: {str(e)}", exc_info=True)
        yield json.dumps({"status": "Failed", "detail": f"Failed during indexing: {str(e)}"}) + "\n"
        
        # ATOMIC ROLLBACK
        logger.warning(f"Starting atomic rollback for {filename} ({doc_id})")
        
        # Rollback metadata from MongoDB
        if metadata_saved:
            try:
                await db.indexed_documents.delete_one({"id": doc_id})
                logger.info("Rollback: MongoDB metadata deleted.")
            except Exception as me:
                logger.error(f"Rollback error: failed to delete MongoDB metadata: {me}")
                
        # Rollback ChromaDB vectors
        if vector_ids:
            try:
                # Retrieve actual vector store and delete
                vector_store.delete(ids=vector_ids)
                logger.info(f"Rollback: deleted {len(vector_ids)} vectors from ChromaDB.")
            except Exception as ve:
                logger.error(f"Rollback error: failed to delete ChromaDB vectors: {ve}")
                
        # Rollback PDF from disk
        if saved_on_disk and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info("Rollback: PDF deleted from disk.")
            except Exception as fe:
                logger.error(f"Rollback error: failed to delete PDF file: {fe}")

async def delete_indexed_document(doc_id: str) -> bool:
    """
    Deletes a document from disk, ChromaDB, and MongoDB metadata.
    """
    db = get_database()
    vector_store = get_vector_store()
    
    # 1. Fetch document metadata
    doc = await db.indexed_documents.find_one({"id": doc_id})
    if not doc:
        logger.warning(f"Document {doc_id} not found in database for deletion.")
        return False
        
    filename = doc.get("filename", "unknown")
    vector_ids = doc.get("vector_ids", [])
    file_path = doc.get("filepath")
    
    # 2. Delete vectors from ChromaDB
    # If vector_ids is empty, try fallback by querying Chroma for doc_id
    if not vector_ids:
        try:
            chroma_res = vector_store.get(where={"doc_id": doc_id})
            vector_ids = chroma_res.get("ids", [])
        except Exception as e:
            logger.error(f"Failed to query ChromaDB for fallback deletion of {doc_id}: {e}")
            
    if vector_ids:
        try:
            vector_store.delete(ids=vector_ids)
            logger.info(f"Deleted {len(vector_ids)} vectors from ChromaDB for doc {doc_id}.")
        except Exception as e:
            logger.error(f"Failed to delete ChromaDB vectors for {doc_id}: {e}")
            
    # 3. Delete file from disk
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
            logger.info(f"Deleted physical file {file_path} for doc {doc_id}.")
        except Exception as e:
            logger.error(f"Failed to delete physical file {file_path}: {e}")
            
    # 4. Delete MongoDB metadata
    try:
        await db.indexed_documents.delete_one({"id": doc_id})
        logger.info(f"Deleted MongoDB metadata for doc {doc_id}.")
        clear_response_cache()
        clear_retriever_cache()
    except Exception as e:
        logger.error(f"Failed to delete MongoDB metadata for {doc_id}: {e}")
        
    # Log Admin Activity
    from app.services.admin_service import log_admin_activity
    await log_admin_activity(
        action="Document Deleted",
        username="admin",
        details={"filename": filename, "doc_id": doc_id}
    )
    
    return True

async def reindex_document_generator(doc_id: str) -> AsyncGenerator[str, None]:
    """
    Reindexes an existing PDF document from disk.
    Yields progress status as line-delimited JSON.
    """
    db = get_database()
    vector_store = get_vector_store()
    
    # 1. Fetch document metadata
    doc = await db.indexed_documents.find_one({"id": doc_id})
    if not doc:
        yield json.dumps({"status": "Failed", "detail": f"Document {doc_id} not found."}) + "\n"
        return
        
    filename = doc["filename"]
    file_path = doc["filepath"]
    old_vector_ids = doc.get("vector_ids", [])
    
    yield json.dumps({"status": "Extracting", "progress": 20}) + "\n"
    
    if not file_path or not os.path.exists(file_path):
        yield json.dumps({"status": "Failed", "detail": f"Physical file not found on disk at {file_path}."}) + "\n"
        # Update metadata status to Failed
        await db.indexed_documents.update_one(
            {"id": doc_id},
            {"$set": {"status": "Failed", "error": "Physical file missing"}}
        )
        return
        
    # Track states for fallback
    new_vector_ids: List[str] = []
    
    try:
        # Extract text page-by-page
        pages_data = extract_text_by_page(file_path)
        
        # 2. Chunking
        yield json.dumps({"status": "Chunking", "progress": 40}) + "\n"
        chunks = chunk_pdf_pages(pages_data, filename, doc_id, doc["hash"])
        
        if not chunks:
            yield json.dumps({"status": "Failed", "detail": "The document contains no extractable text."}) + "\n"
            return
            
        # 3. Embedding
        yield json.dumps({"status": "Embedding", "progress": 60}) + "\n"
        new_vector_ids = [f"{doc_id}_chunk_{i}_{int(datetime.now(timezone.utc).timestamp())}" for i in range(len(chunks))]
        
        # 4. Updating Chroma (Insert new vectors)
        yield json.dumps({"status": "Updating Chroma", "progress": 80}) + "\n"
        vector_store.add_documents(chunks, ids=new_vector_ids)
        
        # 5. Delete old vectors from ChromaDB (Clean up previous index)
        if old_vector_ids:
            try:
                vector_store.delete(ids=old_vector_ids)
                logger.info(f"Deleted {len(old_vector_ids)} old vectors from ChromaDB during reindexing of {doc_id}.")
            except Exception as e:
                logger.error(f"Failed to delete old ChromaDB vectors during reindexing of {doc_id}: {e}")
                
        # 6. Saving Metadata
        yield json.dumps({"status": "Saving Metadata", "progress": 95}) + "\n"
        await db.indexed_documents.update_one(
            {"id": doc_id},
            {
                "$set": {
                    "status": "Indexed",
                    "chunks_count": len(chunks),
                    "vector_ids": new_vector_ids,
                    "updated": datetime.now(timezone.utc)
                },
                "$unset": {"error": ""}
            }
        )
        clear_response_cache()
        clear_retriever_cache()
        
        # Log Admin Activity
        from app.services.admin_service import log_admin_activity
        await log_admin_activity(
            action="Document Reindexed",
            username="admin",
            details={"filename": filename, "doc_id": doc_id, "chunks": len(chunks)}
        )
        
        # 7. Completed
        yield json.dumps({
            "status": "Completed", 
            "progress": 100, 
            "doc_id": doc_id,
            "filename": filename,
            "chunks_count": len(chunks)
        }) + "\n"
        
    except Exception as e:
        logger.error(f"Failed to reindex document {filename} ({doc_id}): {str(e)}", exc_info=True)
        yield json.dumps({"status": "Failed", "detail": f"Failed during reindexing: {str(e)}"}) + "\n"
        
        # Rollback newly added vectors if failed
        if new_vector_ids:
            try:
                vector_store.delete(ids=new_vector_ids)
                logger.info(f"Reindex Rollback: deleted new vectors {len(new_vector_ids)} from ChromaDB.")
            except Exception as ve:
                logger.error(f"Reindex Rollback error: failed to delete new ChromaDB vectors: {ve}")
                
        # Update metadata status to Failed
        await db.indexed_documents.update_one(
            {"id": doc_id},
            {"$set": {"status": "Failed", "error": str(e)}}
        )
